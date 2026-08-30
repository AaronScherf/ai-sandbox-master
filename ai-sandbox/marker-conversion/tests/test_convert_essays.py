import os
import tempfile
import unittest
from unittest.mock import patch

from essays.convert_essays import (
    _unescape_markdown,
    build_frontmatter,
    derive_folder_category,
    discover_docx_files,
    process_docx,
)


class TestDiscoverDocxFiles(unittest.TestCase):
    def test_finds_docx_files_sorted(self):
        with tempfile.TemporaryDirectory() as tmp:
            for name in ["b.docx", "a.docx", "notes.txt"]:
                open(os.path.join(tmp, name), "w").close()
            files = discover_docx_files(tmp)
            self.assertEqual([os.path.basename(f) for f in files], ["a.docx", "b.docx"])

    def test_filters_to_one_file(self):
        with tempfile.TemporaryDirectory() as tmp:
            for name in ["a.docx", "b.docx"]:
                open(os.path.join(tmp, name), "w").close()
            files = discover_docx_files(tmp, file_filter="b.docx")
            self.assertEqual([os.path.basename(f) for f in files], ["b.docx"])

    def test_missing_directory_returns_empty_list(self):
        self.assertEqual(discover_docx_files("/no/such/dir"), [])

    def test_ignores_non_docx_files(self):
        with tempfile.TemporaryDirectory() as tmp:
            open(os.path.join(tmp, "notes.txt"), "w").close()
            self.assertEqual(discover_docx_files(tmp), [])

    def test_ignores_word_lock_files(self):
        # Word creates a "~$name.docx" lock file alongside an open document
        # -- not a real essay, would fail to convert if picked up.
        with tempfile.TemporaryDirectory() as tmp:
            open(os.path.join(tmp, "~$essay.docx"), "w").close()
            open(os.path.join(tmp, "essay.docx"), "w").close()
            files = discover_docx_files(tmp)
            self.assertEqual([os.path.basename(f) for f in files], ["essay.docx"])


class TestDeriveFolderCategory(unittest.TestCase):
    def test_returns_immediate_parent_folder_name(self):
        self.assertEqual(
            derive_folder_category("/root/independent-research/notes/application_essays/Harvard.docx"),
            "application_essays",
        )

    def test_loose_file_directly_under_notes(self):
        self.assertEqual(
            derive_folder_category("/root/independent-research/notes/Research Ideas.docx"),
            "notes",
        )


class TestUnescapeMarkdown(unittest.TestCase):
    def test_unescapes_defensively_escaped_punctuation(self):
        self.assertEqual(_unescape_markdown(r"well\-known, the U\.S\.\!"), "well-known, the U.S.!")

    def test_collapses_doubled_backslash_from_a_literal_source_backslash(self):
        self.assertEqual(_unescape_markdown(r"C:\\Users"), r"C:\Users")

    def test_leaves_ordinary_text_unchanged(self):
        self.assertEqual(_unescape_markdown("Ordinary prose, no escaping needed"), "Ordinary prose, no escaping needed")


class TestBuildFrontmatter(unittest.TestCase):
    def test_renders_flat_metadata_as_yaml(self):
        fm = build_frontmatter({"source_docx": "Essay.docx", "word_count": 500})
        self.assertEqual(
            fm,
            '---\nsource_docx: Essay.docx\nword_count: 500\n---\n\n',
        )

    def test_quotes_values_containing_special_characters(self):
        fm = build_frontmatter({"source_docx": "Statement: MIT.docx"})
        self.assertIn('source_docx: "Statement: MIT.docx"', fm)

    def test_renders_empty_list_bare_not_quoted(self):
        # retag.py's write-back (_TAGS_LINE_RE) expects a bare "tags: [...]"
        # line -- a quoted "tags: \"[]\"" would still match the regex but
        # produce invalid YAML once real tags get substituted in.
        fm = build_frontmatter({"tags": []})
        self.assertIn("tags: []", fm)

    def test_renders_populated_list_bare(self):
        fm = build_frontmatter({"tags": ["personal-statement", "phd-application"]})
        self.assertIn("tags: [personal-statement, phd-application]", fm)


class TestProcessDocx(unittest.TestCase):
    def test_writes_markdown_with_frontmatter(self):
        try:
            import docx as _docx_lib  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed; skipping real .docx round-trip test")

        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            docx_path = os.path.join(tmp, "My Essay.docx")
            doc = Document()
            doc.add_heading("My Essay", level=1)
            doc.add_paragraph("This is the body of the essay.")
            doc.save(docx_path)

            output_dir = os.path.join(tmp, "processed_outputs")
            md_path = process_docx(docx_path, output_dir)

            self.assertEqual(md_path, os.path.join(output_dir, "My Essay.md"))
            with open(md_path, encoding="utf-8") as f:
                content = f.read()
            self.assertIn("source_docx: My Essay.docx", content)
            self.assertIn("tags: []", content)
            self.assertIn("# My Essay", content)
            self.assertIn("This is the body of the essay", content)

    def test_skips_indexing_when_no_index_root_or_client_given(self):
        try:
            import docx as _docx_lib  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed; skipping real .docx round-trip test")

        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            docx_path = os.path.join(tmp, "Essay.docx")
            Document().save(docx_path)

            with patch("essays.convert_essays.reconcile_and_write") as mock_reconcile:
                process_docx(docx_path, os.path.join(tmp, "processed_outputs"))
                mock_reconcile.assert_not_called()

    def test_indexes_with_course_derived_from_index_root(self):
        try:
            import docx as _docx_lib  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed; skipping real .docx round-trip test")

        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            index_root = os.path.join(tmp, "research")
            docx_dir = os.path.join(index_root, "independent-research", "notes", "application_essays")
            os.makedirs(docx_dir)
            docx_path = os.path.join(docx_dir, "Essay.docx")
            Document().save(docx_path)

            output_dir = os.path.join(docx_dir, "processed_outputs")
            fake_client = object()

            with patch("essays.convert_essays.reconcile_and_write") as mock_reconcile:
                process_docx(docx_path, output_dir, index_root=index_root, client=fake_client)

                mock_reconcile.assert_called_once()
                _, kwargs = mock_reconcile.call_args
                self.assertEqual(mock_reconcile.call_args[0][0], index_root)
                self.assertEqual(kwargs["course"], "notes")
                self.assertEqual(kwargs["client"], fake_client)
                self.assertEqual(kwargs["folder_category"], "application_essays")

    def test_indexing_failure_does_not_raise_or_block_conversion(self):
        try:
            import docx as _docx_lib  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed; skipping real .docx round-trip test")

        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            docx_path = os.path.join(tmp, "Essay.docx")
            Document().save(docx_path)

            with patch("essays.convert_essays.reconcile_and_write", side_effect=RuntimeError("boom")):
                # Must not raise -- the .md file is already written and complete
                # regardless of what happens to the indexing hook.
                md_path = process_docx(
                    docx_path, os.path.join(tmp, "processed_outputs"),
                    index_root=tmp, client=object(),
                )
                self.assertTrue(os.path.exists(md_path))


if __name__ == "__main__":
    unittest.main()
